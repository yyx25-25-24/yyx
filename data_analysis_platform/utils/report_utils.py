import os
import datetime
from pathlib import Path
import json


def save_markdown_report(query: str, response: str, dataset_key: str, save_dir: str, charts: list[str] | None = None) -> str:
    os.makedirs(save_dir, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{dataset_key}_{timestamp}.md"
    filename = filename.replace(" ", "_").replace("/", "_")
    save_path = os.path.join(save_dir, filename)

    with open(save_path, "w", encoding="utf-8") as f:
        f.write(f"# 分析报告\n\n")
        f.write(f"- 数据集：{dataset_key}\n")
        f.write(f"- 生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"- 查询内容：{query}\n\n")
        f.write("---\n\n")
        f.write("## 分析结果\n\n")
        # 如果 response 是可解析的 JSON 字符串或 dict，尝试格式化为表格
        parsed = None
        if isinstance(response, dict):
            parsed = response
        else:
            try:
                parsed = json.loads(response)
            except Exception:
                parsed = None

        if isinstance(parsed, dict):
            # 若包含多个表格型字段，逐个导出为 markdown 表格
            for k, v in parsed.items():
                if isinstance(v, list):
                    # list of records -> table
                    try:
                        from pandas import DataFrame
                        df = DataFrame(v)
                        f.write(f"### {k}\n")
                        f.write(df.to_markdown(index=False))
                        f.write("\n\n")
                        continue
                    except Exception:
                        pass
                # 其他字段直接写入
                f.write(f"**{k}**:\n{v}\n\n")
        else:
            f.write(response)

        # 嵌入图表（如果提供）
        if charts:
            f.write("## 图表\n\n")
            for c in charts:
                # 如果是 html，则以链接形式嵌入；如果是图片则显示图片
                ext = os.path.splitext(c)[1].lower()
                if ext in (".html",):
                    f.write(f"- 图表文件：[{os.path.basename(c)}]({c})\n")
                else:
                    # 使用 markdown 图片展示
                    f.write(f"![{os.path.basename(c)}]({c})\n\n")
        f.write("\n")

    return save_path


def save_docx_report(query: str, response: str, dataset_key: str, save_dir: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 库，请先安装：pip install python-docx") from exc

    os.makedirs(save_dir, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{dataset_key}_{timestamp}.docx"
    filename = filename.replace(" ", "_").replace("/", "_")
    save_path = os.path.join(save_dir, filename)

    doc = Document()
    doc.add_heading("分析报告", level=1)
    doc.add_paragraph(f"数据集：{dataset_key}")
    doc.add_paragraph(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"查询内容：{query}")
    doc.add_paragraph("")
    doc.add_heading("分析结果", level=2)
    for line in response.splitlines():
        doc.add_paragraph(line)

    doc.save(save_path)
    return save_path


def save_html_report(query: str, response: str, dataset_key: str, save_dir: str, charts: list[str] | None = None) -> str:
    os.makedirs(save_dir, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{dataset_key}_{timestamp}.html"
    filename = filename.replace(" ", "_").replace("/", "_")
    save_path = os.path.join(save_dir, filename)

    html_lines = []
    html_lines.append('<!doctype html>')
    html_lines.append('<html><head><meta charset="utf-8"><title>分析报告</title></head><body>')
    html_lines.append(f'<h1>分析报告</h1>')
    html_lines.append(f'<p><strong>数据集：</strong>{dataset_key}</p>')
    html_lines.append(f'<p><strong>生成时间：</strong>{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>')
    html_lines.append(f'<p><strong>查询内容：</strong>{query}</p>')
    html_lines.append('<hr>')
    html_lines.append('<h2>分析结果</h2>')

    # format response
    if isinstance(response, dict):
        html_lines.append('<div>')
        for k, v in response.items():
            html_lines.append(f'<h3>{k}</h3>')
            html_lines.append(f'<pre>{json.dumps(v, ensure_ascii=False, indent=2)}</pre>')
        html_lines.append('</div>')
    else:
        html_lines.append(f'<pre>{response}</pre>')

    if charts:
        html_lines.append('<h2>图表</h2>')
        for c in charts:
            # if looks like URL or absolute path starting with / serve as is
            if isinstance(c, str) and (c.startswith('http') or c.startswith('/')):
                src = c
            else:
                # assume filesystem path, use basename so served by /charts/ route
                src = os.path.basename(c)
                src = f'/charts/{src}'
            html_lines.append(f'<div><img src="{src}" style="max-width:100%;height:auto;"></div>')

    html_lines.append('</body></html>')

    with open(save_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_lines))

    return save_path


def save_pdf_report(query: str, response: str, dataset_key: str, save_dir: str, charts: list[str] | None = None) -> str:
    # generate intermediate html then convert via pdfkit if available
    html_path = save_html_report(query, response, dataset_key, save_dir, charts=charts)
    pdf_path = os.path.splitext(html_path)[0] + '.pdf'
    try:
        import pdfkit
    except Exception as exc:
        raise RuntimeError('要生成 PDF，请先安装 pdfkit 并配置 wkhtmltopdf。') from exc
    try:
        pdfkit.from_file(html_path, pdf_path)
    except Exception as exc:
        raise RuntimeError(f'PDF 生成失败: {exc}') from exc
    return pdf_path


def save_analysis_report(query: str, response: str, dataset_key: str, save_dir: str, report_format: str) -> str:
    if report_format.lower() == "md":
        # backward-compatible: accept optional charts passed via tuple
        charts = None
        if isinstance(response, tuple) and len(response) == 2:
            resp_text, charts = response
            return save_markdown_report(query, resp_text, dataset_key, save_dir, charts=charts)
        return save_markdown_report(query, response, dataset_key, save_dir)
    if report_format.lower() == "docx":
        return save_docx_report(query, response, dataset_key, save_dir)
    if report_format.lower() == "html":
        charts = None
        if isinstance(response, tuple) and len(response) == 2:
            resp_text, charts = response
            return save_html_report(query, resp_text, dataset_key, save_dir, charts=charts)
        return save_html_report(query, response, dataset_key, save_dir)
    if report_format.lower() == "pdf":
        charts = None
        if isinstance(response, tuple) and len(response) == 2:
            resp_text, charts = response
            return save_pdf_report(query, resp_text, dataset_key, save_dir, charts=charts)
        return save_pdf_report(query, response, dataset_key, save_dir)
    raise ValueError(f"不支持的报告格式: {report_format}")
